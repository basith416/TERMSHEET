# app.py (product + per-product template + leg-based payoff composition)
import streamlit as st
from docx import Document
import io
import datetime

st.set_page_config(page_title="Term Sheet Generator", layout="wide")
st.title("Term Sheet Generator — Word template + Scenario Table (product-aware)")

# -------------------- helper functions (same as your earlier version) --------------------

def replace_text_in_paragraphs_full(doc, placeholder, replacement):
    for para in doc.paragraphs:
        if placeholder in para.text:
            new_text = para.text.replace(placeholder, str(replacement))
            para.clear()
            para.add_run(new_text)

def replace_text_in_cell_paragraphs_full(cell, placeholder, replacement):
    for para in cell.paragraphs:
        if placeholder in para.text:
            new_text = para.text.replace(placeholder, str(replacement))
            para.clear()
            para.add_run(new_text)

def replace_text_in_tables_full(doc, placeholder, replacement):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_cell_paragraphs_full(cell, placeholder, replacement)

def replace_in_headers_and_footers(doc, placeholder, replacement):
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for para in header.paragraphs:
            if placeholder in para.text:
                new_text = para.text.replace(placeholder, str(replacement))
                para.clear()
                para.add_run(new_text)
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text_in_cell_paragraphs_full(cell, placeholder, replacement)
        for para in footer.paragraphs:
            if placeholder in para.text:
                new_text = para.text.replace(placeholder, str(replacement))
                para.clear()
                para.add_run(new_text)
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text_in_cell_paragraphs_full(cell, placeholder, replacement)

def find_paragraph_with_placeholder(doc, placeholder):
    for para in doc.paragraphs:
        if placeholder in para.text:
            return para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        return para
    for section in doc.sections:
        for para in section.header.paragraphs:
            if placeholder in para.text:
                return para
        for para in section.footer.paragraphs:
            if placeholder in para.text:
                return para
    return None

def insert_table_after_paragraph(doc, paragraph, data, col_names=None, preferred_style_name="Table Grid"):
    nrows = len(data) + (1 if col_names else 0)
    ncols = len(data[0]) if data else (len(col_names) if col_names else 1)
    table = doc.add_table(rows=nrows, cols=ncols)
    try:
        table.style = preferred_style_name
    except Exception:
        pass
    row_idx = 0
    if col_names:
        hdr_cells = table.rows[row_idx].cells
        for c, name in enumerate(col_names):
            hdr_cells[c].text = str(name)
        row_idx += 1
    for r, row in enumerate(data):
        cells = table.rows[row_idx + r].cells
        for c, val in enumerate(row):
            cells[c].text = str(val)
    try:
        paragraph._p.addnext(table._tbl)
    except Exception:
        st.warning("Could not insert table at the placeholder location; table appended at the end.")

# -------------------- UI & inputs --------------------

st.markdown(
    "Upload a Word template (.docx) per product (or use one template for all). "
    "Placeholders supported: `{{ClientName}}`, `{{ValuationDate}}`, `{{MaturityDate}}`, `{{Spot}}`, `{{Strike}}`, `{{Premium}}`, "
    "`{{Notional}}`, `{{ImpliedVol}}`, `{{ScenarioTable}}`."
)

st.sidebar.header("Product & Template")
product = st.sidebar.selectbox("Product", ["Call", "Put", "Range Forward", "Seagull", "I-Forward", "Custom"])

st.sidebar.markdown("Upload a template for the selected product (optional). If not provided, upload a general template below.")
template_for_product = st.sidebar.file_uploader(f"Template for {product} (optional)", type=["docx"], key=f"template_{product}")

st.sidebar.markdown("---")
st.sidebar.markdown("Or upload a single general template (used if product-specific is not provided):")
general_template = st.sidebar.file_uploader("General template (.docx)", type=["docx"], key="template_general")

st.sidebar.markdown("---")
st.sidebar.header("Trade & scenario inputs")
client_name = st.sidebar.text_input("Client name", value="ACME Ltd.")
valuation_date = st.sidebar.date_input("Valuation date", value=datetime.date.today())
maturity_date = st.sidebar.date_input("Maturity date", value=(datetime.date.today() + datetime.timedelta(days=30)))
strike = st.sidebar.number_input("Primary Strike (for convenience)", value=80.0, step=0.5, format="%.2f")
spot = st.sidebar.number_input("Spot (current)", value=78.0, step=0.5, format="%.2f")
premium = st.sidebar.number_input("Premium", value=0.0, step=0.01, format="%.4f")
notional = st.sidebar.number_input("Notional", value=1000000.0, step=1000.0, format="%.2f")
implied_vol = st.sidebar.number_input("Implied vol (in %)", value=12.0, step=0.1, format="%.2f")

st.sidebar.markdown("---")
st.sidebar.header("Scenario grid")
min_spot = st.sidebar.number_input("Scenario: min spot", value=max(0.0, strike - 40.0), step=0.5, format="%.2f")
max_spot = st.sidebar.number_input("Scenario: max spot", value=strike + 40.0, step=0.5, format="%.2f")
step_spot = st.sidebar.number_input("Scenario step", value=2.0, step=0.1, format="%.2f")

# -------------------- leg composition (flexible product definition) --------------------

st.header("Product legs (build your product by adding legs)")
st.markdown(
    "Each leg can be a Call, Put or Forward. Use the multiplier to model long (+1) or short (-1) notional scaling.\n\n"
    "Examples:\n"
    "- **Seagull** can be constructed by combining a Put leg and Call spread legs.\n"
    "- **Range Forward** can be modelled with two forwards or forwards + options.\n"
    "- **I-Forward** is typically a forward with a cap/floor leg — build it by combining legs."
)

max_legs = 4
legs = []
with st.expander("Define legs (up to 4)", expanded=True):
    n_legs = st.number_input("Number of legs", min_value=1, max_value=max_legs, value=2, step=1)
    for i in range(int(n_legs)):
        st.markdown(f"**Leg {i+1}**")
        leg_type = st.selectbox(f"Type (leg {i+1})", ["Call", "Put", "Forward"], key=f"leg_type_{i}")
        leg_strike = st.number_input(f"Strike (leg {i+1})", value=float(strike), key=f"leg_strike_{i}", format="%.2f")
        leg_mult = st.number_input(f"Multiplier / notional factor (leg {i+1})", value=1.0, step=0.1, key=f"leg_mult_{i}", format="%.2f")
        legs.append({"type": leg_type, "strike": float(leg_strike), "mult": float(leg_mult)})

# -------------------- template selection logic --------------------

# choose the product template if provided, else fall back to general template
template_file = template_for_product if template_for_product is not None else general_template

if not template_file:
    st.warning("Upload at least one Word template (.docx) — either product-specific or general — to enable document generation.")
    st.stop()

# -------------------- preview --------------------

st.write("### Preview inputs")
st.write(f"**Client**: {client_name}  •  **Product**: {product}  •  **Strike**: {strike}  •  **Spot**: {spot}")
st.write(f"Valuation date: {valuation_date}  •  Maturity date: {maturity_date}")
st.write(f"Premium: {premium}  •  Notional: {notional}  •  Implied vol: {implied_vol}%")
st.write(f"Scenario from {min_spot} to {max_spot} step {step_spot}")
st.write("Legs:")
for i, leg in enumerate(legs):
    st.write(f"- Leg {i+1}: {leg['type']} @ {leg['strike']:.2f} × {leg['mult']:.2f}")

# -------------------- payoff calculators --------------------

def payoff_call(spot_val, K):
    return max(0.0, spot_val - K)

def payoff_put(spot_val, K):
    return max(0.0, K - spot_val)

def payoff_forward(spot_val, K):
    # linear forward payoff (spot - strike)
    return spot_val - K

def combined_payoff(spot_val, legs_list):
    total = 0.0
    for lg in legs_list:
        if lg["type"] == "Call":
            total += lg["mult"] * payoff_call(spot_val, lg["strike"])
        elif lg["type"] == "Put":
            total += lg["mult"] * payoff_put(spot_val, lg["strike"])
        elif lg["type"] == "Forward":
            total += lg["mult"] * payoff_forward(spot_val, lg["strike"])
    return round(total, 2)

# -------------------- Generate term sheet button --------------------

if st.button("Generate Term Sheet"):

    # load document
    try:
        template_doc = Document(template_file)
    except Exception as e:
        st.error(f"Failed to read the uploaded docx template: {e}")
        st.stop()

    # scenario spots
    spots = []
    val = float(min_spot)
    max_sp = float(max_spot)
    s_step = float(step_spot) if step_spot > 0 else 1.0
    while val <= max_sp + 1e-9:
        spots.append(round(val, 2))
        val += s_step

    # build scenario table rows: Spot at maturity, Payoff (sum of legs), and an optional scaled P&L by notional
    scenario_rows = []
    for s in spots:
        payoff = combined_payoff(s, legs)
        # if you prefer scaled by notional multiply here; we include both unscaled and scaled
        payoff_scaled = round(payoff * (notional / 1.0), 2)
        scenario_rows.append([s, payoff, payoff_scaled])

    # placeholders formatting
    placeholders = {
        "{{ClientName}}": client_name,
        "{{ValuationDate}}": valuation_date.strftime("%Y-%m-%d"),
        "{{MaturityDate}}": maturity_date.strftime("%Y-%m-%d"),
        "{{Strike}}": "{:.2f}".format(strike),
        "{{Spot}}": "{:.2f}".format(spot),
        "{{Premium}}": "{:.4f}".format(premium),
        "{{Notional}}": "{:,.2f}".format(notional),
        "{{ImpliedVol}}": "{:.2f}%".format(implied_vol),
        "{{Product}}": product,
    }

    # replace placeholders everywhere
    for ph, val in placeholders.items():
        replace_text_in_paragraphs_full(template_doc, ph, val)
        replace_text_in_tables_full(template_doc, ph, val)
        replace_in_headers_and_footers(template_doc, ph, val)

    # insert scenario table at placeholder or append
    placeholder = "{{ScenarioTable}}"
    para = find_paragraph_with_placeholder(template_doc, placeholder)
    col_names = ["Spot at Maturity", "Payoff (per unit)", f"Payoff × Notional ({notional:,.2f})"]
    if para:
        if placeholder in para.text:
            new_text = para.text.replace(placeholder, "")
            para.clear()
            para.add_run(new_text)
        insert_table_after_paragraph(template_doc, para, scenario_rows, col_names=col_names)
    else:
        insert_table_after_paragraph(template_doc, template_doc.paragraphs[-1], scenario_rows, col_names=col_names)
        st.info("Placeholder {{ScenarioTable}} not found — scenario table appended at document end.")

    # Save and provide download
    output = io.BytesIO()
    template_doc.save(output)
    output.seek(0)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_client = client_name.replace(" ", "_")
    filename = f"TermSheet_{product}_{safe_client}_{timestamp}.docx"

    st.success("Document ready — click the button below to download.")
    st.download_button(
        label="Download Term Sheet (.docx)",
        data=output.getvalue(),
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.info("Click **Generate Term Sheet** to create and download the document.")
